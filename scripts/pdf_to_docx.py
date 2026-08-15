#!/usr/bin/env python3
"""
Wrapper CLI tipis di atas pdf2docx utk dipanggil sbg subprocess dari Node.js
(services/convertService.js). Sengaja dibuat sesederhana mungkin: terima path
input & output, jalankan konversi, exit code 0 = sukses, non-zero = gagal
(pesan error di stderr, bisa dibaca Node utk fallback ke jalur JS).

Setelah konversi utama, ditambahkan post-processing opsional: PDF sering
punya link internal (mis. entri Daftar Isi -> heading terkait) yang TIDAK
direkonstruksi pdf2docx (cuma menangani hyperlink URI eksternal). Kalau ada,
link-link ini diekstrak lewat PyMuPDF lalu dipasang ulang sbg bookmark +
hyperlink internal Word. Ini enhancement, BUKAN bagian kritis -- kalau
gagal/error, konversi utama tetap dianggap berhasil (skip diam-diam, cuma
warning ke stderr), krn dokumen tanpa link internal masih jauh lebih baik
drpd tidak ada dokumen sama sekali.

Usage: python3 pdf_to_docx.py <input.pdf> <output.docx>
"""
import sys
import os
import re
import difflib


def strip_toc_label(src_text):
    """'Register......................5' -> 'Register'. Entri Daftar Isi PDF
    biasanya diakhiri deretan titik (dot leader) lalu nomor halaman -- ini
    dibuang supaya tersisa cuma judul section yg mau dicocokkan ke heading
    aslinya di halaman tujuan."""
    m = re.match(r'^(.*?)[.\s]{3,}\d*\s*$', src_text)
    return (m.group(1) if m else src_text).strip()


def get_page_text_blocks(page):
    """Ambil semua blok teks di satu halaman PDF, tiap blok jadi satu string
    (gabungan semua span di dalamnya), zero-width-space dibuang krn kadang
    dipakai PDF sbg pemisah 'tak terlihat' antara dua heading yg berdekatan
    tanpa jeda baris sungguhan (mis. 'BAB II' + ZWSP + 'PEMBAHASAN')."""
    blocks = page.get_text('dict')['blocks']
    out = []
    for b in blocks:
        if b.get('type') != 0:
            continue
        text = ''.join(s['text'] for l in b['lines'] for s in l['spans'])
        text = text.replace('\u200b', '').strip()
        if text:
            out.append(text)
    return out


def find_best_text_match(blocks, label):
    """3 lapis strategi cocokkan label Daftar Isi ke teks di halaman tujuan:
    1) exact match (kasus paling umum)
    2) startswith (utk blok gabungan spt 'BAB IIPEMBAHASAN')
    3) fuzzy match, ambang 0.75 (utk typo kecil di dokumen sumber, mis.
       'Desktop' di Daftar Isi vs 'Dekstop' di heading aslinya -- ditemukan
       nyata di salah satu dokumen uji, bukan kasus teoretis)
    Return None kalau tidak ada yg cukup yakin -- lebih baik skip drpd salah
    pasang link ke tempat yang salah."""
    label_norm = label.lower().strip()
    if not label_norm:
        return None
    for b in blocks:
        if b.lower().strip() == label_norm:
            return b
    for b in blocks:
        if len(label_norm) >= 3 and b.lower().strip().startswith(label_norm):
            return b
    best, best_ratio = None, 0
    for b in blocks:
        ratio = difflib.SequenceMatcher(None, b.lower().strip(), label_norm).ratio()
        if ratio > best_ratio:
            best_ratio, best = ratio, b
    return best if best_ratio >= 0.75 else None


def extract_internal_links(pdf_path):
    """Ekstrak semua link internal (GOTO, bukan URI eksternal) dari SETIAP
    halaman PDF, resolve teks sumber (label di halaman asal) & teks target
    (heading di halaman tujuan). Return list of {label, target_text}."""
    import fitz  # pylint: disable=import-outside-toplevel

    doc = fitz.open(pdf_path)
    results = []
    try:
        for page in doc:
            for link in page.get_links():
                if link.get('kind') != fitz.LINK_GOTO and link.get('kind') != 4:
                    continue  # cuma link internal, lewati URI eksternal (sudah ditangani pdf2docx)
                target_page_idx = link.get('page')
                target_point = link.get('to')
                if target_page_idx is None or target_point is None:
                    continue  # destinasi tidak ke-resolve di level PDF -- lewati, tidak bisa diperbaiki dari sini
                try:
                    src_text = page.get_textbox(link['from']).strip()
                except Exception:  # pylint: disable=broad-except
                    continue
                label = strip_toc_label(src_text)
                if not label:
                    continue
                target_blocks = get_page_text_blocks(doc[target_page_idx])
                target_text = find_best_text_match(target_blocks, label)
                if target_text:
                    results.append({'label': label, 'target_text': target_text})
    finally:
        doc.close()
    return results


def add_internal_hyperlinks(docx_path, links):
    """Pasang bookmark di paragraf/run yang cocok dgn target_text, lalu bungkus
    run yang cocok dgn label sbg hyperlink internal ke bookmark itu. Kerja di
    level RUN (bukan paragraf) krn pdf2docx sering menggabungkan banyak baris
    PDF (mis. seluruh Daftar Isi) jadi SATU paragraf dgn satu run per baris."""
    import docx  # pylint: disable=import-outside-toplevel
    from docx.oxml.ns import qn  # pylint: disable=import-outside-toplevel
    from docx.oxml import OxmlElement  # pylint: disable=import-outside-toplevel

    document = docx.Document(docx_path)
    all_paragraphs = document.paragraphs

    def make_bookmark(bookmark_id, name):
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), str(bookmark_id))
        start.set(qn('w:name'), name)
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), str(bookmark_id))
        return start, end

    def wrap_run_as_hyperlink(run_element, anchor_name):
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), anchor_name)
        hyperlink.set(qn('w:history'), '1')
        parent = run_element.getparent()
        parent.insert(list(parent).index(run_element), hyperlink)
        parent.remove(run_element)
        hyperlink.append(run_element)

    bookmark_id = 100  # mulai dari angka besar, hindari bentrok id internal Word lain
    applied = 0
    used_target_paragraphs = set()  # id(paragraph) yg sudah dipakai -- cegah heading yg
    # teksnya berulang (mis. 'Register' muncul lagi utk section Desktop stlh Mobile)
    # selalu ke-bookmark ke kemunculan PERTAMA utk SEMUA link yg labelnya sama.
    for link in links:
        # 1) Cari paragraf target (persis/mengandung target_text) yg BELUM dipakai
        #    link lain & pasang bookmark -- urutan link mengikuti urutan halaman PDF,
        #    jadi maju ke kemunculan berikutnya scr alami utk heading yg berulang.
        target_para = None
        for p in all_paragraphs:
            if id(p) in used_target_paragraphs:
                continue
            if p.text.strip() == link['target_text'].strip() or p.text.strip().startswith(link['target_text'].strip()):
                target_para = p
                break
        if target_para is None or not target_para.runs:
            continue
        used_target_paragraphs.add(id(target_para))

        anchor_name = f"pdflink_{bookmark_id}"
        start_el, end_el = make_bookmark(bookmark_id, anchor_name)
        first_run_el = target_para.runs[0]._r  # pylint: disable=protected-access
        first_run_el.addprevious(start_el)
        first_run_el.addnext(end_el)
        bookmark_id += 1

        # 2) Cari run sumber (persis label ATAU diawali label -- run TOC biasa
        #    berisi label + dot leader + nomor halaman dalam satu run) & bungkus jadi hyperlink
        wrapped = False
        for p in all_paragraphs:
            if wrapped:
                break
            for r in p.runs:
                r_text = r.text.strip()
                if r_text.lower() == link['label'].strip().lower() or r_text.lower().startswith(link['label'].strip().lower() + '.'):
                    wrap_run_as_hyperlink(r._r, anchor_name)  # pylint: disable=protected-access
                    wrapped = True
                    applied += 1
                    break

    document.save(docx_path)
    return applied


def fix_toc_line_breaks(docx_path):
    """PDF sering merender Daftar Isi dgn dot leader (banyak titik) yg secara
    visual pas krn PDF adalah format LAYOUT TETAP (fixed-width per halaman).
    pdf2docx merekonstruksi teks dot leader itu APA ADANYA sbg karakter titik
    literal yg sangat panjang (bisa 100+ titik per entri), TANPA tab stop
    dot-leader asli Word DAN tanpa line break antar entri -- hasilnya satu
    paragraf raksasa (bisa 3000+ karakter) yg di Word (format REFLOW, lebar
    bisa berubah-ubah) di-word-wrap seenaknya di titik manapun, bikin
    tampilan berantakan (deretan titik itu sendiri kepanjangan utk satu
    baris, DAN antar entri saling menyambung tanpa jeda baris).

    Fix (2 bagian, keduanya perlu supaya benar-benar rapi):
    1. Tambahkan tab stop RIGHT-ALIGNED dgn leader titik ke pPr paragraf --
       Word yg akan menggambar titik-titiknya secara dinamis pas ke lebar
       kolom yg tersedia, bukan teks statis yg bisa kepanjangan/kependekan.
    2. Ganti teks 'Label.......N' tiap entri jadi 'Label' + TAB + 'N', dan
       sisipkan <w:br/> line break sesudahnya supaya tiap entri mulai di
       baris barunya sendiri."""
    import docx  # pylint: disable=import-outside-toplevel
    from docx.oxml.ns import qn  # pylint: disable=import-outside-toplevel
    from docx.oxml import OxmlElement  # pylint: disable=import-outside-toplevel

    document = docx.Document(docx_path)
    W_T = qn('w:t')
    # '....5' atau '.... 5 ' dst -- leader titik (min 3 titik berturutan)
    # diikuti nomor halaman di akhir teks elemen.
    entry_pattern = re.compile(r'(\.{3,})\s*(\d+)\s*$')
    fixed_paragraphs = 0

    for paragraph in document.paragraphs:
        p_element = paragraph._p  # pylint: disable=protected-access
        full_text = ''.join(t.text or '' for t in p_element.iter(W_T))
        if '.' * 15 not in full_text or len(full_text) < 200:
            continue  # bukan paragraf Daftar Isi (heuristik: dot leader panjang & teks total panjang)

        # Ambil lebar halaman & indentasi paragraf ini utk hitung posisi tab
        # yg pas (dekat tepi kanan area teks yg tersedia utk paragraf ini).
        sect_pr = p_element.getroottree().getroot().find(f'.//{qn("w:sectPr")}')
        pg_sz = sect_pr.find(qn('w:pgSz')) if sect_pr is not None else None
        page_width = int(pg_sz.get(qn('w:w'))) if pg_sz is not None else 11920
        ind_el = p_element.find(f'{qn("w:pPr")}/{qn("w:ind")}')
        right_indent = int(ind_el.get(qn('w:right'), '0')) if ind_el is not None else 0
        tab_pos = max(1000, page_width - right_indent - 200)  # -200 = buffer utk digit nomor halaman

        # 1) Pasang tab stop dot-leader ke pPr paragraf (cuma sekali per paragraf)
        p_pr = p_element.find(qn('w:pPr'))
        if p_pr is None:
            p_pr = OxmlElement('w:pPr')
            p_element.insert(0, p_pr)
        tabs_el = p_pr.find(qn('w:tabs'))
        if tabs_el is None:
            tabs_el = OxmlElement('w:tabs')
            p_pr.append(tabs_el)
        tab_def = OxmlElement('w:tab')
        tab_def.set(qn('w:val'), 'right')
        tab_def.set(qn('w:leader'), 'dot')
        tab_def.set(qn('w:pos'), str(tab_pos))
        tabs_el.append(tab_def)

        # 2) Ganti tiap 'Label.......N' jadi 'Label' + TAB + 'N', sisipkan
        # line break sesudahnya.
        children = list(p_element)
        touched_any = False
        for child in children:
            tag = child.tag.split('}')[-1]
            if tag not in ('r', 'hyperlink'):
                continue
            t_elements = list(child.iter(W_T))
            if not t_elements:
                continue
            last_t = t_elements[-1]
            text = last_t.text or ''
            m = entry_pattern.search(text)
            if not m:
                continue
            # Buang leader titik, sisakan label (kalau ada di elemen yg sama)
            # + tab + nomor halaman.
            prefix = text[:m.start()]
            page_num = m.group(2)
            last_t.text = prefix
            last_t.set(qn('xml:space'), 'preserve')

            tab_run = OxmlElement('w:r')
            tab_char = OxmlElement('w:tab')
            tab_run.append(tab_char)
            num_t = OxmlElement('w:t')
            num_t.text = page_num
            num_t.set(qn('xml:space'), 'preserve')
            tab_run.append(num_t)
            child.addnext(tab_run)

            br_run = OxmlElement('w:r')
            br_run.append(OxmlElement('w:br'))
            tab_run.addnext(br_run)
            touched_any = True

        if touched_any:
            fixed_paragraphs += 1

    if fixed_paragraphs:
        document.save(docx_path)
    return fixed_paragraphs


def fix_duplicated_table_columns(docx_path):
    """Bug spesifik ditemukan di pdf2docx: utk tabel tanpa garis pemisah
    terlihat di PDF sumber (kolom disusun cuma berdasarkan alignment posisi,
    bukan grid asli), kadang SEMUA nilai satu kolom (dipisah newline) malah
    ke-duplikasi APA ADANYA di SETIAP baris, bukan didistribusikan satu nilai
    per baris spt seharusnya (mis. tabel nama+NIM: kolom NIM di semua baris
    isinya SAMA PERSIS -- keempat NIM digabung newline -- padahal harusnya
    baris 1 cuma NIM orang pertama, baris 2 NIM orang kedua, dst).

    Root cause lengkapnya: pdf2docx jg keliru memasang vMerge (merge sel
    vertikal) di kolom itu -- kadang pola restart/continue-nya sendiri tidak
    konsisten antar baris (diverifikasi lgs dari XML). Karena python-docx
    mengembalikan OBJEK SEL YANG SAMA utk semua baris yg ter-vMerge, cuma
    redistribusi teks TANPA membongkar vMerge dulu bikin semua baris balik
    ke-overwrite jadi nilai yg sama (baris terakhir yg diproses menang).

    Fix: lepas SEMUA vMerge di tabel yg kena masalah ini dulu (bikin tiap
    sel independen), baru redistribusi teksnya satu nilai per baris."""
    import docx  # pylint: disable=import-outside-toplevel
    from docx.oxml.ns import qn  # pylint: disable=import-outside-toplevel

    document = docx.Document(docx_path)
    fixed_columns = 0

    for table in document.tables:
        n_rows = len(table.rows)
        if n_rows < 2:
            continue
        n_cols = len(table.columns)

        # Deteksi dulu kolom mana yg kena pola bug (semua sel identik &
        # jumlah baris teks cocok jumlah baris tabel) SEBELUM ubah apapun,
        # krn setelah vMerge dilepas nanti cell.text scr sementara jadi tidak
        # bisa diandalkan (python-docx butuh tabel dlm keadaan konsisten).
        buggy_cols = []
        for col_idx in range(n_cols):
            cell_texts = [table.rows[r].cells[col_idx].text for r in range(n_rows)]
            first = cell_texts[0].strip()
            if not first or any(t.strip() != first for t in cell_texts):
                continue
            lines = [ln.strip() for ln in first.split('\n') if ln.strip()]
            if len(lines) == n_rows:
                buggy_cols.append((col_idx, lines))

        if not buggy_cols:
            continue

        # Lepas SEMUA vMerge di tabel ini -- lebih aman drpd cuma di kolom
        # yg buggy, krn strukturnya sendiri terbukti bisa tidak konsisten
        # antar kolom (vMerge kolom lain bisa saja masih merujuk baris yg
        # sama scr XML meski logicalnya beda kolom).
        for tc in table._tbl.iter(qn('w:tc')):  # pylint: disable=protected-access
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is None:
                continue
            v_merge = tc_pr.find(qn('w:vMerge'))
            if v_merge is not None:
                tc_pr.remove(v_merge)

        for col_idx, lines in buggy_cols:
            # Ambil referensi formatting dari run PERTAMA yg ketemu di kolom
            # ini (biasanya baris 0, sel "restart" dari vMerge lama) supaya
            # baris lain yg run-nya kosong (sel lanjutan vMerge, formatting-
            # nya ikut hilang) tetap konsisten drpd jatuh ke default Word.
            ref_size = ref_bold = ref_font = None
            for r in range(n_rows):
                existing_runs = table.rows[r].cells[col_idx].paragraphs[0].runs if table.rows[r].cells[col_idx].paragraphs else []
                if existing_runs and existing_runs[0].font.size:
                    ref_size = existing_runs[0].font.size
                    ref_bold = existing_runs[0].font.bold
                    ref_font = existing_runs[0].font.name
                    break

            for r in range(n_rows):
                cell = table.rows[r].cells[col_idx]
                paragraphs = cell.paragraphs
                if not paragraphs:
                    continue
                for run in list(paragraphs[0].runs):
                    run.text = ''
                if paragraphs[0].runs:
                    target_run = paragraphs[0].runs[0]
                    target_run.text = lines[r]
                else:
                    target_run = paragraphs[0].add_run(lines[r])
                if not target_run.font.size and ref_size:
                    target_run.font.size = ref_size
                    target_run.font.bold = ref_bold
                    target_run.font.name = ref_font
                for extra_p in paragraphs[1:]:
                    for run in extra_p.runs:
                        run.text = ''
            fixed_columns += 1

    if fixed_columns:
        document.save(docx_path)
    return fixed_columns


def postprocess_table_duplication(docx_path):
    """Enhancement non-kritis, pola sama spt post-processing lainnya."""
    try:
        fixed = fix_duplicated_table_columns(docx_path)
        if fixed:
            print(f'Kolom tabel dgn nilai terduplikasi diperbaiki: {fixed}.', file=sys.stderr)
    except Exception as e:  # pylint: disable=broad-except
        print(f'Post-processing perbaikan tabel dilewati (non-fatal): {type(e).__name__}: {e}', file=sys.stderr)


def postprocess_toc_formatting(docx_path):
    """Enhancement non-kritis, pola sama spt postprocess_internal_links --
    gagal di sini TIDAK menggagalkan konversi utama."""
    try:
        fixed = fix_toc_line_breaks(docx_path)
        if fixed:
            print(f'Line break Daftar Isi diperbaiki di {fixed} paragraf.', file=sys.stderr)
    except Exception as e:  # pylint: disable=broad-except
        print(f'Post-processing format Daftar Isi dilewati (non-fatal): {type(e).__name__}: {e}', file=sys.stderr)


def postprocess_internal_links(pdf_path, docx_path):
    """Enhancement non-kritis -- gagal di sini TIDAK menggagalkan konversi
    utama, cuma di-skip diam-diam (warning ke stderr saja)."""
    try:
        links = extract_internal_links(pdf_path)
        if not links:
            return
        applied = add_internal_hyperlinks(docx_path, links)
        print(f'Link internal terpasang: {applied}/{len(links)}', file=sys.stderr)
    except Exception as e:  # pylint: disable=broad-except
        print(f'Post-processing link internal dilewati (non-fatal): {type(e).__name__}: {e}', file=sys.stderr)


def main():
    if len(sys.argv) != 3:
        print('Usage: pdf_to_docx.py <input.pdf> <output.docx>', file=sys.stderr)
        sys.exit(2)

    input_path, output_path = sys.argv[1], sys.argv[2]

    if not os.path.isfile(input_path):
        print(f'File input tidak ditemukan: {input_path}', file=sys.stderr)
        sys.exit(1)

    try:
        # Import di dalam try -- kalau pdf2docx/PyMuPDF tidak ter-install
        # (mis. environment dev lokal tanpa Docker), ini akan gagal dgn
        # ImportError yang jelas, ditangkap Node sbg sinyal utk fallback ke
        # konversi berbasis pdfjs-dist yang sudah ada, bukan crash total.
        from pdf2docx import Converter
    except ImportError as e:
        print(f'pdf2docx tidak tersedia: {e}', file=sys.stderr)
        sys.exit(3)

    try:
        cv = Converter(input_path)
        try:
            cv.convert(output_path)
        finally:
            cv.close()
    except Exception as e:  # pylint: disable=broad-except
        print(f'Konversi gagal: {type(e).__name__}: {e}', file=sys.stderr)
        sys.exit(4)

    if not os.path.isfile(output_path) or os.path.getsize(output_path) == 0:
        print('Output tidak dihasilkan atau kosong.', file=sys.stderr)
        sys.exit(5)

    postprocess_internal_links(input_path, output_path)
    postprocess_toc_formatting(output_path)
    postprocess_table_duplication(output_path)

    print(f'OK: {output_path} ({os.path.getsize(output_path)} bytes)')
    sys.exit(0)


if __name__ == '__main__':
    main()
